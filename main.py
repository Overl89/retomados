from __future__ import print_function
import datetime
import pandas
from pandas import *
import pyodbc

from mailmerge import MailMerge
from docx.shared import Inches
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from sqlalchemy import create_engine


from selenium.webdriver.common.keys import Keys
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.action_chains import ActionChains
import pandas as pd
import openpyxl
import time
import json
import csv
import json
from datetime import date
import sys


# driver = webdriver.Firefox()
# driver.get("https://safedataanalytics.com.br/")
# driver.maximize_window()
# delay = 3
#
# login = driver.find_element(By.CSS_SELECTOR, "div.form-group:nth-child(1)>input:nth-child(2)")
# login.send_keys('lucivando.santos@sistemastecnol.com.br')
#
# time.sleep(60)
# hoje = date.today()
# datafinal=hoje - datetime.timedelta(days=365)
# #api = driver.find_element(By.CSS_SELECTOR, "li.dropdown:nth-child(2)>span:nth-child(1)")
# #api.click()
# dia=datafinal.day
# print(hoje)
# print(datafinal)
#
# time.sleep(3)
#
# df = pd.read_excel('C:/Users/lucivando.santos/OneDrive - TECNOL SISTEMAS DE AUTOMAÇÂO S A/GPI COMUM/Gestão de processos RPA/Robô Retomados/Multas detalhadas/Placas.xlsx',usecols=['Placa'])
# dataframe = openpyxl.load_workbook('C:/Users/lucivando.santos/OneDrive - TECNOL SISTEMAS DE AUTOMAÇÂO S A/GPI COMUM/Gestão de processos RPA/Robô Retomados/Multas detalhadas/Placas.xlsx')
# dataframe1 = dataframe.active
# linhas = len(df)
# #Pesquisa Multas Detalhadas
# #multas = driver.find_element(By.CSS_SELECTOR, "ul.show>li:nth-child(4)>a:nth-child(1)")
# #multas.click()
# new_url = "https://safedataanalytics.com.br/api/veiculo/multasdetalhadas"
# driver.execute_script("window.open('');")
# driver.switch_to.window(driver.window_handles[1])
# driver.get(new_url)
# #
# # for row in dataframe1.iter_rows(min_row=2, min_col=1, max_row=linhas + 1, max_col=1):
# #      for cell in row:
# #          placa= driver.find_element(By.CSS_SELECTOR, "div.col-lg-3:nth-child(1)>div:nth-child(1)>input:nth-child(2)")
# #          placa.clear()
# #          placa.send_keys(cell.value)
# #          time.sleep(2)
# #          dtinicial = driver.find_element(By.CSS_SELECTOR, "div.col-lg-3:nth-child(2)>div:nth-child(1)>input:nth-child(2)")
# #          time.sleep(2)
# #          ActionChains(driver).move_to_element(dtinicial).click().send_keys("{:02d}".format(datafinal.month) + "{:02d}".format(datafinal.day) +str(datafinal.year)).perform()
# #          time.sleep(2)
# #          dtfinal = driver.find_element(By.CSS_SELECTOR, "div.col-lg-3:nth-child(3)>div:nth-child(1)>input:nth-child(2)")
# #          ActionChains(driver).move_to_element(dtfinal).click().send_keys("{:02d}".format(hoje.month) + "{:02d}".format(hoje.day) + str(hoje.year)).perform()
# #          time.sleep(1)
# #          pesquisar=driver.find_element(By.CSS_SELECTOR, ".btn")
# #          pesquisar.click()
# #          time.sleep(5)
# #          json=driver.find_element(By.CSS_SELECTOR, "#multas-panel>div:nth-child(1)>div:nth-child(1)>pre:nth-child(1)")
# #          arquivo= json.get_property('innerHTML')
# #          f = open("C:/Users/lucivando.santos/OneDrive - TECNOL SISTEMAS DE AUTOMAÇÂO S A/GPI COMUM/Gestão de processos RPA/Robô Retomados/Multas detalhadas/Resultados"+"/"+cell.value+"_MultasDetalhadas.txt", "w")
# #          f.write(arquivo)
# #          f.close()
# #          json_str = arquivo
# #          df = pd.read_json(json_str)
# #          df.to_csv('courses.csv')
# #          df.to_json
# #
# # driver.close()
#
# # Switching to old tab
# driver.switch_to.window(driver.window_handles[0])
#
# new_url = "https://safedataanalytics.com.br/api/veiculo/detalhada"
# driver.execute_script("window.open('');")
# driver.switch_to.window(driver.window_handles[1])
# driver.get(new_url)
#
#
# #api = driver.find_element(By.CSS_SELECTOR, "li.dropdown:nth-child(2)>span:nth-child(1)")
# #api.click()
#
# #veidetalhada = driver.find_element(By.CSS_SELECTOR, "ul.show>li:nth-child(1)>a:nth-child(1)")
# #veidetalhada.click()
# # pesquisa desabilitada devido duplicidade de dados
# #pesquisa Veiculos detalhado
# #for row in dataframe1.iter_rows(min_row=2, min_col=1, max_row=linhas + 1, max_col=1):
#      #for cell in row:
#         #placa=driver.find_element(By.CSS_SELECTOR, ".form-control")
#         #placa.clear()
#         #placa.send_keys(cell.value)
#         #time.sleep(2)
#         #pesquisa = driver.find_element(By.CSS_SELECTOR, ".btn>span:nth-child(1)")
#         #pesquisa.click()
#         #time.sleep(5)
#         #json=driver.find_element(By.CSS_SELECTOR, "#detalhes-panel>div:nth-child(1)>div:nth-child(1)>pre:nth-child(1)")
#         #arquivo = json.get_property('innerHTML')
#         #f = open("C:/Users/lucivando.santos/OneDrive - TECNOL SISTEMAS DE AUTOMAÇÂO S A/GPI COMUM/Gestão de processos RPA/Robô Retômados/Veículo Detalhado/Resultados" + "/" + cell.value + "_VeiDetalhada.txt","w")
#         #f.write(arquivo)
#         #f.close()
#
# #driver.close()
#
#
# # Switching to old tab
# # driver.switch_to.window(driver.window_handles[0])
# #
# # new_url = "https://safedataanalytics.com.br/api/veiculo/fiscalizacao"
# # driver.execute_script("window.open('');")
# # driver.switch_to.window(driver.window_handles[1])
# # driver.get(new_url)
# # #api = driver.find_element(By.CSS_SELECTOR, "li.dropdown:nth-child(2)>span:nth-child(1)")
# # #api.click()
# #
# # #fiscalizacao = driver.find_element(By.CSS_SELECTOR, "ul.show>li:nth-child(2)>a:nth-child(1)")
# # #fiscalizacao.click()
# # #Pesquisa Fiscalização
# # for row in dataframe1.iter_rows(min_row=2, min_col=1, max_row=linhas + 1, max_col=1):
# #      for cell in row:
# #         placa = driver.find_element(By.CSS_SELECTOR, ".form-control")
# #         placa.clear()
# #         placa.send_keys(cell.value)
# #         time.sleep(2)
# #         pesquisa = driver.find_element(By.CSS_SELECTOR, ".btn>span:nth-child(1)")
# #         pesquisa.click()
# #         time.sleep(5)
# #         json = driver.find_element(By.CSS_SELECTOR,"#restricoes-panel>div:nth-child(1)>div:nth-child(1)>pre:nth-child(1)")
# #         arquivo = json.get_property('innerHTML')
# #         f = open("C:/Users/lucivando.santos/OneDrive - TECNOL SISTEMAS DE AUTOMAÇÂO S A/GPI COMUM/Gestão de processos RPA/Robô Retomados/Fiscalização/Resultados" + "/" + cell.value + "_Fiscalizacao.txt","w")
# #         f.write(arquivo)
# #         f.close()
# #
# # driver.close()
#
# # Switching to old tab
# driver.switch_to.window(driver.window_handles[0])
# #Pesquisa Precificação
# #api = driver.find_element(By.CSS_SELECTOR, "li.dropdown:nth-child(2)>span:nth-child(1)")
# #api.click()
#
# #precificacao = driver.find_element(By.CSS_SELECTOR, "li.dropdown:nth-child(2)>ul:nth-child(2)>li:nth-child(5)>a:nth-child(1)")
# #precificacao.click()
#
# new_url = "https://safedataanalytics.com.br/api/veiculo/precificacao"
# driver.execute_script("window.open('');")
# driver.switch_to.window(driver.window_handles[1])
# driver.get(new_url)
#
# for row in dataframe1.iter_rows(min_row=2, min_col=1, max_row=linhas + 1, max_col=1):
#      for cell in row:
#         placa = driver.find_element(By.CSS_SELECTOR, ".form-control")
#         placa.clear()
#         placa.send_keys(cell.value)
#         time.sleep(2)
#         pesquisa = driver.find_element(By.CSS_SELECTOR, ".btn>span:nth-child(1)")
#         pesquisa.click()
#         time.sleep(5)
#         json = driver.find_element(By.CSS_SELECTOR,"#restricoes-panel>div:nth-child(1)>div:nth-child(1)>pre:nth-child(1)")
#         arquivo = json.get_property('innerHTML')
#         f = open("C:/Users/lucivando.santos/OneDrive - TECNOL SISTEMAS DE AUTOMAÇÂO S A/GPI COMUM/Gestão de processos RPA/Robô Retomados/Precificação/Resultados" + "/" + cell.value + "_Precificacao.txt","w")
#         f.write(arquivo)
#         f.close()
#
# driver.close()
#
# # Switching to old tab
# driver.switch_to.window(driver.window_handles[0])
# driver.close()


placa='PYZ1I63'



# Opening JSON file
with open('C:/Users/lucivando.santos/OneDrive - TECNOL SISTEMAS DE AUTOMAÇÂO S A/GPI COMUM/Gestão de processos RPA/Robô Retomados/Precificação/Resultados/Processados/'+ placa +'_Precificacao.txt') as json_file:
   data = json.load(json_file)
   json_str = json.dumps(data)


   resp = json.loads(json_str)

   modelo=resp['data']['result']['modelo']
   placa=resp['data']['placa_pesquisada']
   valor=resp['data']['result']['valorFIPE']

   template = "C:/Users/lucivando.santos/OneDrive - TECNOL SISTEMAS DE AUTOMAÇÂO S A/GPI COMUM/Gestão de processos RPA/Robô Retomados/MascaraDossieLocalizacao.docx"


document = MailMerge(template)


document.merge(Modelo=modelo,ValorFipe=valor,Placa=placa)

with open('C:/Users/lucivando.santos/OneDrive - TECNOL SISTEMAS DE AUTOMAÇÂO S A/GPI COMUM/Gestão de processos RPA/Robô Retomados/Fiscalização/Resultados/Processados/'+ placa +'_Fiscalizacao.txt') as json_file:
   data = json.load(json_file)
   json_str = json.dumps(data)
   resp = json.loads(json_str)

   proprietario=resp['data']['result']['veiculo'][0]['nomeProprietario']
   possuidor=resp['data']['result']['veiculo'][0]['possuidor']['nome']
   chassi=resp['data']['result']['veiculo'][0]['restricoes']['restricao'][0]['chassi']
   #
document.merge(Possuidor=possuidor,Proprietario=proprietario,Chassi=chassi)

with open('C:/Users/lucivando.santos/OneDrive - TECNOL SISTEMAS DE AUTOMAÇÂO S A/GPI COMUM/Gestão de processos RPA/Robô Retomados/Fiscalização/Resultados/Processados/'+ placa +'_Fiscalizacao.txt') as json_file:
   data = json.load(json_file)
   json_str = json.dumps(data)
   resp = json.loads(json_str)


   restricao=resp['data']['result']['veiculo'][0]['restricoes']['restricao'][0]['tipoRestricao']
   dtrestricao=resp['data']['result']['veiculo'][0]['restricoes']['restricao'][0]['dataHoraAtualizacao']
   dt=dtrestricao[0:10]
   dtAno=dt[0:4]
   dtmes=dt[5:7]
   dtdia=dt[8:10]
   document.merge(DataRestricao=dtdia +'/'+ dtmes + '/'+ dtAno, DescRestricao=restricao)


with open('C:/Users/lucivando.santos/OneDrive - TECNOL SISTEMAS DE AUTOMAÇÂO S A/GPI COMUM/Gestão de processos RPA/Robô Retomados/Multas detalhadas/Resultados/Processados/'+ placa +'_MultasDetalhadas.txt') as json_file:
   data = json.load(json_file)
   json_str = json.dumps(data)
   resp = json.loads(json_str)


   dtmulta=resp['data']['infracoes'][0]['dataInfracao']
   Descmulta=resp['data']['infracoes'][0]['descricaoInfracao']
   valormulta=resp['data']['infracoes'][0]['valorIntegralInfracao']
   horamulta=resp['data']['infracoes'][0]['horaInfracao']

   dt = dtmulta[0:10]
   dtAno = dt[0:4]
   dtmes = dt[5:7]
   dtdia = dt[8:10]
   #pd.json_normalize(data['data']['infracoes']).to_csv('jsonMultasJson.csv')
   multas=pd.json_normalize(data['data']['infracoes']).to_string()




   #document.merge(DtMulta=dtdia +'/'+ dtmes + '/'+ dtAno+ ' ' + horamulta,DescMulta=Descmulta,ValorMulta=valormulta)


   #driver = webdriver.Firefox()
   #driver.get("https://www.google.com/")
   #driver.maximize_window()

   #pesquisa = driver.find_element(By.CSS_SELECTOR, ".gLFyf")
   #pesquisa.send_keys('rua Ary fernandes 31,São Paulo')
   #pesquisa.send_keys(Keys.ENTER)

time.sleep(1)

#enter= driver.find_element(By.CSS_SELECTOR,"div.lJ9FBc:nth-child(9)>center:nth-child(2)>input:nth-child(1)")
#enter.click()


#image=driver.find_element(By.CSS_SELECTOR,"#lu_map")
#image.screenshot('C:/Users/lucivando.santos/OneDrive - TECNOL SISTEMAS DE AUTOMAÇÂO S A/GPI COMUM/Gestão de processos RPA/Robô Retômados/QNM7A61/localizacao.png')

#driver.close()

conn = pyodbc.connect(r'Driver={Microsoft Access Driver (*.mdb, *.accdb)};DBQ=C:\Users\lucivando.santos\OneDrive - TECNOL SISTEMAS DE AUTOMAÇÂO S A\GPI COMUM\Gestão de processos RPA\Robô Retômados\BaseRetomados.accdb;')
cursor = conn.cursor()
cursor.execute('Select DataTexto,Endereço,TipoRegistro From TbRodagem where Placa =? ORDER BY Data DESC',placa)
df=pd.DataFrame(cursor.fetchall())

df[['DtLocalizacao','Endlocalizacao','Detalhe']] = df[0].apply(lambda x: pd.Series(str(x).split(",")))
df.__delitem__(0)
doc = Document('C:/Users/lucivando.santos/OneDrive - TECNOL SISTEMAS DE AUTOMAÇÂO S A/GPI COMUM/Gestão de processos RPA/Robô Retomados/DossieTecshare_teste.docx')

# add a table to the end and create a reference variable
# extra row is so we can add the header row
t = doc.add_table(rows=df.shape[0],cols=df.shape[1])


string=df.to_string()

df[['DtLocalizacao', 'Endlocalizacao','Detalhe']].to_excel(r'df.xlsx')
excel=pandas.read_excel('df.xlsx')

df.to_json
json_str=excel.to_json(orient='records')
json_str=json_str.replace('(','').replace(')','')
columns = json_str.replace("\\u00a0", "")
columns = json_str.replace("()", "")
columns = json_str.replace("'", "")
columns = json.dumps(columns)
columns = json.loads(columns)
array = '{"columns": %s}' % columns
data = json.loads(array)


document.merge_rows('DtLocalizacao',data['columns'])

#for i in range(df.shape[0]):
#  for j in range(df.shape[1]):
# cell = df.iat[i, j]
#  t.cell(i, j).text = str(cell)



#document.write('C:/Users/lucivando.santos/OneDrive - TECNOL SISTEMAS DE AUTOMAÇÂO S A/GPI COMUM/Gestão de processos RPA/Robô Retômados/DossieTecshare_teste.docx')

tipoop="Multa"
data=cursor.execute('SELECT DataTexto,descricaoInfracao, valor FROM TbRodagem WHERE (((TipoRegistro)=?) AND ((placa)=?))  ORDER BY Data DESC',tipoop,placa)
df=pd.DataFrame(cursor.fetchall())
df[['DtMulta', 'DescMulta','ValorMulta']] = df[0].apply(lambda x: pd.Series(str(x).split(",")))

df[['DtMulta','DescMulta','ValorMulta']].to_excel(r'Multas.xlsx')
excel=pandas.read_excel('Multas.xlsx')


df.to_json
json_str=excel.to_json(orient='records')
json_str=json_str.replace("00'))",'').replace('Decimal(','').replace("'",'').replace('(','').replace('.',',')
columns = json_str
columns = json.dumps(columns)
columns = json.loads(columns)
columns
array = '{"columns": %s}' % columns
data = json.loads(array)

document.merge_rows('DtMulta',data['columns'])


data=cursor.execute('SELECT Cliente, UF FROM TbCliente WHERE PLACA=?;',placa)
df=pd.DataFrame(cursor.fetchall())
df[['Cliente', 'UF',]] = df[0].apply(lambda x: pd.Series(str(x).split(",")))

document.merge(Cliente=df['Cliente'].to_string().replace('(','').replace("'","").replace('0    ',''),UfCtt=df['UF'].to_string().replace(')','').replace("'","").replace('0     ',''))

document.write('C:/Users/lucivando.santos/OneDrive - TECNOL SISTEMAS DE AUTOMAÇÂO S A/GPI COMUM/Gestão de processos RPA/Robô Retomados/DossieTecshare_'+ placa +'.docx')

document1 = Document('C:/Users/lucivando.santos/OneDrive - TECNOL SISTEMAS DE AUTOMAÇÂO S A/GPI COMUM/Gestão de processos RPA/Robô Retomados/DossieTecshare_'+ placa +'.docx')
image_paras = [i for i, p in enumerate(document1.paragraphs) if "LOCALIZAÇÃO" in p.text]
p = document1.paragraphs[image_paras[0]]
p.text = ""
r = p.add_run()
r.add_picture('C:/Users/lucivando.santos/OneDrive - TECNOL SISTEMAS DE AUTOMAÇÂO S A/GPI COMUM/Gestão de processos RPA/Robô Retomados/Fotos/'+ placa +'.jpg',width=Inches(3),height=Inches(2))
r.add_picture('C:/Users/lucivando.santos/OneDrive - TECNOL SISTEMAS DE AUTOMAÇÂO S A/GPI COMUM/Gestão de processos RPA/Robô Retomados/QNM7A61/ImagemUltimoLocal.png',width=Inches(3),height=Inches(2))
document1.save('C:/Users/lucivando.santos/OneDrive - TECNOL SISTEMAS DE AUTOMAÇÂO S A/GPI COMUM/Gestão de processos RPA/Robô Retomados/Dossie/DossieTecshare_'+ placa +'.docx')

